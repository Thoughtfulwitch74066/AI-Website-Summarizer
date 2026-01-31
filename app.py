from flask import Flask, render_template, request, send_file
import requests
from bs4 import BeautifulSoup
import re
import os
import time
import logging
from groq import Groq
from dotenv import load_dotenv
from urllib.parse import urlparse
from docx import Document

load_dotenv()

app = Flask(__name__)

# -----------------------------
# Logging Configuration
# -----------------------------

logging.basicConfig(
    filename="app.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# -----------------------------
# Utility Functions
# -----------------------------

def is_valid_url(url):
    parsed = urlparse(url)
    return parsed.scheme in ("http", "https") and parsed.netloc != ""


def fetch_website_content(url):
    headers = {"User-Agent": "Mozilla/5.0"}

    try:
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        raise Exception(f"Failed to fetch website: {str(e)}")

    soup = BeautifulSoup(response.text, "html.parser")

    # Remove unwanted tags
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    main_content = soup.find("article") or soup.find("main") or soup.body

    if not main_content:
        raise Exception("Could not extract main content.")

    text_elements = main_content.find_all(["h1", "h2", "h3", "p", "li"])

    content = ""
    for element in text_elements:
        text = element.get_text().strip()
        if text:
            content += text + "\n\n"

    content = re.sub(r"\[\d+\]", "", content)
    content = re.sub(r"\n\s*\n", "\n\n", content)

    if len(content) < 200:
        raise Exception("Extracted content too small. Page may not contain article text.")

    return content.strip()


def get_instruction(summary_type):
    if summary_type == "short":
        return "Provide a short summary in 5-6 lines."
    elif summary_type == "detailed":
        return "Provide a detailed summary explaining key concepts."
    elif summary_type == "bullets":
        return "Provide a bullet-point summary."
    return "Summarize clearly."


def summarize_text(text, summary_type, model_name):

    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise Exception("GROQ_API_KEY not configured.")

    client = Groq(api_key=api_key)

    instruction = get_instruction(summary_type)
    prompt = f"{instruction}\n\n{text[:4000]}"

    start_time = time.time()

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5,
    )

    end_time = time.time()
    duration = round(end_time - start_time, 2)

    summary = response.choices[0].message.content

    return summary, duration


def save_summary_as_word(summary, model_name, summary_type, url):
    document = Document()
    document.add_heading("AI Website Summary", level=1)
    document.add_paragraph(f"Source URL: {url}")
    document.add_paragraph(f"Model Used: {model_name}")
    document.add_paragraph(f"Summary Type: {summary_type}")
    document.add_paragraph(f"Generated At: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("\n")
    document.add_paragraph(summary)

    filename = f"summary_{int(time.time())}.docx"
    document.save(filename)

    return filename


# -----------------------------
# Routes
# -----------------------------

@app.route("/", methods=["GET", "POST"])
def index():
    summary = None
    error = None
    response_time = None
    download_file = None

    if request.method == "POST":
        url = request.form.get("url")
        summary_type = request.form.get("summary_type")
        model_name = request.form.get("model_name")

        if not is_valid_url(url):
            error = "Invalid URL. Please enter a valid http/https URL."
            logging.warning("Invalid URL entered.")
        else:
            try:
                content = fetch_website_content(url)

                summary, response_time = summarize_text(
                    content,
                    summary_type,
                    model_name
                )

                download_file = save_summary_as_word(
                    summary,
                    model_name,
                    summary_type,
                    url
                )

                logging.info(f"Summary generated successfully using model {model_name}")

            except Exception as e:
                error = f"Error: {str(e)}"
                logging.error(str(e))

    return render_template(
        "index.html",
        summary=summary,
        error=error,
        response_time=response_time,
        download_file=download_file
    )


@app.route("/download/<filename>")
def download(filename):
    return send_file(filename, as_attachment=True)


# -----------------------------
# Production Server Config
# -----------------------------

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
